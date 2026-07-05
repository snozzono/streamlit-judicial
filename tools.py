"""
tools.py — Herramientas del agente tributario.

Seis funciones que el grafo LangGraph invoca como nodos o directamente:
  1. buscar_normativa        — FAISS vectorstore normativa
  2. buscar_casos_anteriores — FAISS índice de largo plazo
  3. evaluar_consulta        — LLM evalúa suficiencia del contexto
  4. redactar_memo           — genera .docx en memos/
  5. guardar_drive           — sube a Google Drive vía MCP (opcional)
  6. enviar_gmail            — envía por Gmail vía MCP (opcional)
"""

import importlib
import json
import logging
import os
import time
from datetime import datetime

from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from langchain_core.messages import HumanMessage, SystemMessage
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from openai import APIError, APIConnectionError, APITimeoutError, RateLimitError

from config import CONFIG
from memory import get_memoria_largo_plazo

logger = logging.getLogger(__name__)


def _api_key() -> str:
    return os.getenv("GH_TOKEN") or os.getenv("GITHUB_TOKEN") or ""


def _get_llm() -> ChatOpenAI:
    return ChatOpenAI(
        model=CONFIG.llm_model,
        base_url=CONFIG.api_base_url,
        api_key=_api_key(),
        temperature=0.1,
    )


def _get_embeddings() -> OpenAIEmbeddings:
    return OpenAIEmbeddings(
        model=CONFIG.embedding_model,
        base_url=CONFIG.api_base_url,
        api_key=_api_key(),
    )


# ---------------------------------------------------------------------------
# Retry wrapper con backoff exponencial
# ---------------------------------------------------------------------------

def llamar_con_reintento(llm_func, *args, max_attempts: int = None, **kwargs):
    """
    Envuelve llamadas a LLM con reintento automático y backoff exponencial.
    Solo reintenta ante errores transitorios de API (RateLimitError, APITimeoutError,
    APIConnectionError, APIError). Otros errores se lanzan de inmediato.

    Args:
        llm_func: función a llamar (ej. llm.invoke)
        max_attempts: máximo de intentos (default CONFIG.retry_max_attempts)

    Returns:
        resultado de llm_func

    Raises:
        última excepción después de agotar intentos
    """
    max_attempts = max_attempts or CONFIG.retry_max_attempts
    errores_transitorios = (RateLimitError, APITimeoutError, APIConnectionError, APIError)
    last_error = None
    for intento in range(1, max_attempts + 1):
        try:
            return llm_func(*args, **kwargs)
        except errores_transitorios as e:
            last_error = e
            error_name = type(e).__name__
            if intento < max_attempts:
                delay = CONFIG.retry_initial_delay * (CONFIG.retry_backoff_factor ** (intento - 1))
                logger.warning(
                    f"Intento {intento}/{max_attempts} falló ({error_name}): {e}. "
                    f"Reintentando en {delay:.1f}s..."
                )
                time.sleep(delay)
            else:
                logger.error(
                    f"Intento {intento}/{max_attempts} falló ({error_name}): {e}. "
                    "Sin reintentos restantes."
                )
    raise last_error


# ---------------------------------------------------------------------------
# Clasificador de complejidad para K dinámico
# ---------------------------------------------------------------------------

def estimar_k_optimo(consulta: str) -> int:
    """
    Estima el valor óptimo de K según la complejidad de la consulta.

    Consultas cortas/específicas → K bajo (5-8)
    Consultas largas/complejas  → K alto (10-20)

    Returns:
        entero entre CONFIG.k_min y CONFIG.k_max
    """
    longitud = len(consulta)
    palabras_clave_complejas = [
        "comparar", "diferencia", "relación", "versus", "vs",
        "todos", "cada", "complete", "detalle", "exhaustivo",
        "transversal", "integral", "marco", "general",
    ]
    palabras_clave_especificas = [
        "artículo", "art.", "circular", "resolución",
        "número", "num", "literal", "inciso",
    ]

    # Base según longitud
    if longitud < 50:
        k_base = CONFIG.k_min
    elif longitud < 150:
        k_base = CONFIG.k_default_dynamic
    else:
        k_base = CONFIG.k_min + 5

    consulta_lower = consulta.lower()

    # Penalizar por palabras clave complejas
    for kw in palabras_clave_complejas:
        if kw in consulta_lower:
            k_base += 3
            break

    # Reducir por palabras clave específicas
    for kw in palabras_clave_especificas:
        if kw in consulta_lower:
            k_base -= 1
            break

    return max(CONFIG.k_min, min(CONFIG.k_max, k_base))


# ---------------------------------------------------------------------------
# 1. buscar_normativa
# ---------------------------------------------------------------------------

def buscar_normativa(query: str, k: int = None) -> list[Document]:
    """
    Busca los k fragmentos más relevantes en el vectorstore de normativa.
    
    Si k no se especifica, usa estimación dinámica según la consulta.
    """
    if k is None:
        k = estimar_k_optimo(query)
    try:
        embeddings = _get_embeddings()
        vs = FAISS.load_local(
            CONFIG.vectorstore_dir,
            embeddings,
            allow_dangerous_deserialization=True,
        )
        docs = vs.similarity_search(query, k=k)
        logger.info(f"[tools] buscar_normativa: k={k}, docs={len(docs)} para query='{query[:60]}...'")
        return docs
    except Exception as e:
        logger.error(f"[tools] buscar_normativa: {e}")
        return []


# ---------------------------------------------------------------------------
# 2. buscar_casos_anteriores
# ---------------------------------------------------------------------------

def buscar_casos_anteriores(query: str, k: int = 3) -> list[Document]:
    """Recupera los k casos más similares del índice de largo plazo."""
    return get_memoria_largo_plazo().buscar_casos_similares(query, k=k)


# ---------------------------------------------------------------------------
# 3. evaluar_consulta
# ---------------------------------------------------------------------------

_SYSTEM_EVAL = """Eres un evaluador de calidad de un sistema RAG tributario chileno.
Determina si el contexto normativo recuperado es suficiente para responder la consulta.

Responde SOLO con JSON válido (sin markdown) con estas claves:
- "cubierta": true | false — si la consulta puede responderse con el contexto
- "confianza": float 0.0–1.0 — qué tan bien el contexto cubre la consulta
- "accion_sugerida": "responder" | "buscar_mas" | "redactar_memo"
"""


def evaluar_consulta(consulta: str, contexto: str) -> dict:
    """
    Evalúa si el contexto recuperado cubre la consulta.

    Returns:
        dict con claves: cubierta (bool), confianza (float), accion_sugerida (str)
    """
    from cache import get_cache
    cache = get_cache()
    
    # Intentar desde caché
    cache_key = f"eval:{consulta.strip().lower()}"
    cached = cache.obtener(cache_key)
    if cached:
        try:
            return json.loads(cached)
        except Exception:
            pass

    llm = _get_llm()
    user_msg = (
        f"Consulta: {consulta}\n\n"
        f"Contexto disponible (primeros 2 000 caracteres):\n{contexto[:2000]}\n\n"
        "Evalúa y responde en JSON."
    )
    try:
        resp = llamar_con_reintento(
            llm.invoke,
            [SystemMessage(content=_SYSTEM_EVAL), HumanMessage(content=user_msg)],
        )
        contenido = resp.content.strip().removeprefix("```json").removesuffix("```").strip()
        resultado = json.loads(contenido)

        # Guardar en caché
        try:
            cache.guardar(cache_key, json.dumps(resultado))
        except Exception:
            pass

        return resultado
    except Exception as e:
        logger.error(f"[tools] evaluar_consulta: {e}")
        return {"cubierta": True, "confianza": 0.8, "accion_sugerida": "responder"}


# ---------------------------------------------------------------------------
# 4. redactar_memo
# ---------------------------------------------------------------------------

def redactar_memo(caso: str, contexto: str, analisis: str, destinatario: str) -> str:
    """
    Genera un memorándum tributario formal en .docx.

    Args:
        caso:         descripción del caso / consulta original
        contexto:     normativa relevante recuperada
        analisis:     análisis jurídico generado por el agente
        destinatario: nombre del destinatario del memo

    Returns:
        ruta absoluta del archivo .docx generado
    """
    from docx import Document as DocxDoc
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    os.makedirs(CONFIG.memos_dir, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    ruta = os.path.join(CONFIG.memos_dir, f"memo_{timestamp}.docx")

    doc = DocxDoc()

    titulo = doc.add_heading("MEMORÁNDUM TRIBUTARIO", 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    fecha_str = datetime.now().strftime("%d de %B de %Y")
    doc.add_paragraph(f"Fecha:   {fecha_str}")
    doc.add_paragraph(f"Para:    {destinatario}")
    doc.add_paragraph("De:      Asistente Tributario — Bufete Ruiz Salazar Tributaria")
    doc.add_paragraph("Asunto:  Análisis de Consulta Tributaria")
    doc.add_paragraph()

    doc.add_heading("I. Antecedentes", level=1)
    doc.add_paragraph(caso)

    doc.add_heading("II. Normativa Aplicable", level=1)
    doc.add_paragraph(contexto[:1500] if contexto else "Ver análisis.")

    doc.add_heading("III. Análisis y Conclusiones", level=1)
    doc.add_paragraph(analisis)

    doc.add_heading("IV. Disclaimer", level=1)
    doc.add_paragraph(
        "Este memorándum tiene carácter orientativo y no constituye asesoría legal vinculante. "
        "Se recomienda validar sus conclusiones con un contador o abogado tributario habilitado."
    )

    doc.save(ruta)
    return ruta


# ---------------------------------------------------------------------------
# 5. guardar_drive  — retorna los bytes del archivo para descarga en browser
# ---------------------------------------------------------------------------

def guardar_drive(ruta: str) -> dict:
    """
    Devuelve los bytes del .docx para que la UI ofrezca descarga directa.
    (Google Drive requiere OAuth; para Streamlit local se usa st.download_button.)
    """
    try:
        with open(ruta, "rb") as f:
            datos = f.read()
        nombre = os.path.basename(ruta)
        return {"ok": True, "mensaje": "Archivo listo para descargar.", "ruta": ruta, "bytes": datos, "nombre": nombre}
    except Exception as e:
        return {"ok": False, "mensaje": f"No se pudo leer el archivo: {e}", "ruta": ruta}


# ---------------------------------------------------------------------------
# 6. enviar_gmail  — SMTP con credenciales en .env
#
# Variables requeridas en .env:
#   SMTP_SERVER=smtp.gmail.com
#   SMTP_PORT=587
#   SMTP_USER=tu@gmail.com
#   SMTP_PASS=xxxx xxxx xxxx xxxx   ← App Password de Google (16 chars)
# ---------------------------------------------------------------------------

def enviar_gmail(ruta: str, destinatario: str) -> dict:
    """Envía el memo como adjunto usando SMTP con credenciales del .env."""
    import smtplib
    from email.message import EmailMessage

    servidor = os.getenv("SMTP_SERVER", "smtp.gmail.com")
    puerto   = int(os.getenv("SMTP_PORT", "587"))
    usuario  = os.getenv("SMTP_USER", "")
    password = os.getenv("SMTP_PASS", "")

    if not usuario or not password:
        return {
            "ok": False,
            "mensaje": (
                "Configura SMTP_USER y SMTP_PASS en tu archivo .env. "
                "Para Gmail crea un App Password en myaccount.google.com → Seguridad → Contraseñas de aplicación."
            ),
            "ruta": ruta,
        }

    try:
        with open(ruta, "rb") as f:
            datos_adjunto = f.read()
        nombre_archivo = os.path.basename(ruta)

        msg = EmailMessage()
        msg["Subject"] = "Memorándum Tributario — Ruiz Salazar Tributaria"
        msg["From"]    = usuario
        msg["To"]      = destinatario
        msg.set_content(
            "Estimado/a cliente,\n\n"
            "Adjunto encontrará el memorándum tributario generado por el asistente.\n\n"
            "Saludos,\nBufete Ruiz Salazar Tributaria"
        )
        msg.add_attachment(
            datos_adjunto,
            maintype="application",
            subtype="vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=nombre_archivo,
        )

        with smtplib.SMTP(servidor, puerto) as smtp:
            smtp.starttls()
            smtp.login(usuario, password)
            smtp.send_message(msg)

        return {"ok": True, "mensaje": f"Memo enviado a {destinatario}.", "ruta": ruta}

    except Exception as e:
        return {"ok": False, "mensaje": f"Error al enviar email: {e}", "ruta": ruta}